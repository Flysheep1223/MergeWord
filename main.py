import os
from tkinter import Label, Button, Entry, filedialog, messagebox, Listbox, END, Toplevel, EXTENDED, Frame

import ttkbootstrap as ttk
from docx import Document
from docxcompose.composer import Composer


def merge_doc(source_file_path_list, target_file_path):
    try:
        # 填充分页符文档
        page_break_doc = Document()
        page_break_doc.add_page_break()
        # 定义新文档
        target_doc = Document(source_file_path_list[0])
        target_composer = Composer(target_doc)
        for i in range(len(source_file_path_list)):
            # 跳过第一个作为模板的文件
            if i == 0:
                continue
            # 填充分页符文档
            target_composer.append(page_break_doc)
            # 拼接文档内容
            f = source_file_path_list[i]
            target_composer.append(Document(f))
        # 保存目标文档
        target_composer.save(target_file_path)
        messagebox.showinfo("Success", "Documents merged successfully!")
    except Exception as e:
        messagebox.showerror("Error", str(e))


def select_source_files():
    files = filedialog.askopenfilenames(filetypes=[("Word documents", "*.docx")])
    if files:
        open_sort_window(files)


def open_sort_window(files):
    sort_window = Toplevel(root)
    sort_window.title("Sort Files")
    sort_window.geometry("1200x800")  # 设置窗口大小

    main_frame = Frame(sort_window)
    main_frame.pack(fill='both', expand=True, padx=10, pady=10)

    listbox_frame = Frame(main_frame)
    listbox_frame.pack(side='left', fill='both', expand=True)

    button_frame = Frame(main_frame)
    button_frame.pack(side='right', fill='y')

    listbox = Listbox(listbox_frame, selectmode=EXTENDED)
    listbox.pack(fill='both', expand=True)

    for file in files:
        listbox.insert(END, file)

    def move_up():
        selection = listbox.curselection()
        if not selection:
            return
        for i in selection:
            if i == 0:
                continue
            text = listbox.get(i)
            listbox.delete(i)
            listbox.insert(i - 1, text)
            listbox.selection_set(i - 1)

    def move_down():
        selection = listbox.curselection()
        if not selection:
            return
        for i in reversed(selection):
            if i == listbox.size() - 1:
                continue
            text = listbox.get(i)
            listbox.delete(i)
            listbox.insert(i + 1, text)
            listbox.selection_set(i + 1)

    def confirm_selection():
        sorted_files = listbox.get(0, END)
        source_files_entry.delete(0, 'end')
        source_files_entry.insert(0, ';'.join(sorted_files))
        sort_window.destroy()

    Button(button_frame, text="Move Up", command=move_up).pack(fill='x', pady=5)
    Button(button_frame, text="Move Down", command=move_down).pack(fill='x', pady=5)
    Button(button_frame, text="Confirm", command=confirm_selection).pack(fill='x', pady=5)


def select_target_file():
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])
    target_file_entry.delete(0, 'end')
    target_file_entry.insert(0, file_path)


def merge_files():
    source_files = source_files_entry.get().split(';')
    target_file = target_file_entry.get()

    if not source_files or not target_file:
        messagebox.showwarning("Input required", "Please select both source files and target file.")
        return

    if not all([os.path.isfile(file) for file in source_files]):
        messagebox.showwarning("File error", "One or more source files do not exist.")
        return

    merge_doc(source_files, target_file)


# 创建主窗口
root = ttk.Window(themename="darkly")
root.geometry("1100x300")
root.title("Word Document Merger")

# 创建标签和输入框
source_files_label = Label(root, text="")  # 布局占位符
source_files_label.grid(row=0, column=0, padx=10, pady=10)

source_files_label = Label(root, text="Source Files:")
source_files_label.grid(row=1, column=0, padx=10, pady=10)

source_files_entry = Entry(root, width=50)
source_files_entry.grid(row=1, column=1, padx=10, pady=10)

source_files_button = Button(root, text="Browse...", command=select_source_files)
source_files_button.grid(row=1, column=2, padx=10, pady=10)

target_file_label = Label(root, text="Target File:")
target_file_label.grid(row=2, column=0, padx=10, pady=10)

target_file_entry = Entry(root, width=50)
target_file_entry.grid(row=2, column=1, padx=10, pady=10)

target_file_button = Button(root, text="Browse...", command=select_target_file)
target_file_button.grid(row=2, column=2, padx=10, pady=10)

merge_button = Button(root, text="Merge", command=merge_files)
merge_button.grid(row=3, column=0, columnspan=3, pady=20)

# 运行主循环
root.mainloop()
